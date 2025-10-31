#!/usr/bin/env python3
"""
Cover Letter to Word Document Converter
Converts all Kenneth Davis cover letter text files to professionally formatted Word documents
"""

import os
import glob
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class CoverLetterConverter:
    def __init__(self):
        self.desktop_path = r"C:\Users\kedavis\Desktop"
        self.current_dir = "."
        
    def create_formatted_document(self, content: str, company_name: str) -> Document:
        """Create a professionally formatted Word document from cover letter content."""
        doc = Document()
        
        # Set document margins (1 inch all around)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Split content into lines
        lines = content.strip().split('\n')
        
        # Process each line and apply appropriate formatting
        in_contact_info = True
        in_letter_body = False
        
        for line in lines:
            line = line.strip()
            if not line:  # Skip empty lines except in letter body
                if in_letter_body:
                    doc.add_paragraph()
                continue
            
            # Header/Contact Information
            if in_contact_info and not line.startswith('Dear'):
                if 'Kenneth Davis' in line and 'QA Lead' not in line:
                    # Main name header
                    name_para = doc.add_paragraph()
                    name_run = name_para.add_run(line)
                    name_run.font.size = Pt(16)
                    name_run.bold = True
                    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'QA Lead for Global AI and Automation Strategy' in line:
                    # Title
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(line)
                    title_run.font.size = Pt(12)
                    title_run.italic = True
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'KennethdavisQA@gmail.com' in line or 'LinkedIn:' in line or 'Portfolio:' in line:
                    # Contact information
                    contact_para = doc.add_paragraph()
                    contact_run = contact_para.add_run(line)
                    contact_run.font.size = Pt(11)
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('October') or line.startswith('Hiring Manager') or company_name.replace(' ', '').lower() in line.replace(' ', '').lower():
                    # Date and recipient info
                    doc.add_paragraph()  # Space before date
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(line)
                    date_run.font.size = Pt(11)
                    if 'October' in line:
                        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    # Other header info
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(line)
                    header_run.font.size = Pt(11)
            
            # Start of letter body
            elif line.startswith('Dear'):
                in_contact_info = False
                in_letter_body = True
                doc.add_paragraph()  # Space before Dear
                dear_para = doc.add_paragraph()
                dear_run = dear_para.add_run(line)
                dear_run.font.size = Pt(11)
            
            # Letter body content
            elif in_letter_body:
                if line.startswith('•'):
                    # Bullet points
                    bullet_para = doc.add_paragraph()
                    bullet_run = bullet_para.add_run(line)
                    bullet_run.font.size = Pt(11)
                    bullet_para.style = 'List Bullet'
                elif line.startswith('What I will bring') or line.startswith('Recent achievements') or line.startswith('Recent analytics achievements') or line.startswith('Recent production achievements') or line.startswith('Recent quality management achievements') or line.startswith('Recent enterprise achievements') or line.startswith('Recent live service achievements'):
                    # Section headers
                    doc.add_paragraph()  # Space before section
                    section_para = doc.add_paragraph()
                    section_run = section_para.add_run(line)
                    section_run.font.size = Pt(12)
                    section_run.bold = True
                elif line.startswith('My technical') or line.startswith('My production') or line.startswith('My analytical') or line.startswith('My quality') or line.startswith('My comprehensive') or line.startswith('My enterprise'):
                    # Technical skills paragraph
                    doc.add_paragraph()  # Space before
                    tech_para = doc.add_paragraph()
                    tech_run = tech_para.add_run(line)
                    tech_run.font.size = Pt(11)
                elif line.startswith('I am') and ('excited' in line or 'eager' in line or 'comfortable' in line or 'particularly' in line or 'deeply' in line):
                    # Personal statement paragraphs
                    doc.add_paragraph()  # Space before
                    personal_para = doc.add_paragraph()
                    personal_run = personal_para.add_run(line)
                    personal_run.font.size = Pt(11)
                elif line.startswith('Thank you'):
                    # Closing paragraph
                    doc.add_paragraph()  # Space before
                    closing_para = doc.add_paragraph()
                    closing_run = closing_para.add_run(line)
                    closing_run.font.size = Pt(11)
                elif line.startswith('Sincerely'):
                    # Signature block
                    doc.add_paragraph()  # Space before signature
                    sig_para = doc.add_paragraph()
                    sig_run = sig_para.add_run(line)
                    sig_run.font.size = Pt(11)
                elif line.startswith('Kenneth Davis') and 'QA Lead' in line:
                    # Signature name and title
                    name_para = doc.add_paragraph()
                    name_run = name_para.add_run(line)
                    name_run.font.size = Pt(11)
                    name_run.bold = True
                elif 'KennethdavisQA@gmail.com' in line and 'LinkedIn:' in line:
                    # Contact info in signature
                    contact_para = doc.add_paragraph()
                    contact_run = contact_para.add_run(line)
                    contact_run.font.size = Pt(10)
                elif 'LinkedIn:' in line and 'Portfolio:' in line:
                    # Links in signature
                    links_para = doc.add_paragraph()
                    links_run = links_para.add_run(line)
                    links_run.font.size = Pt(10)
                else:
                    # Regular paragraph content
                    body_para = doc.add_paragraph()
                    body_run = body_para.add_run(line)
                    body_run.font.size = Pt(11)
        
        return doc
    
    def extract_company_name(self, filename: str) -> str:
        """Extract company name from filename."""
        # Remove prefix and suffix to get company name
        name = filename.replace('Kenneth_Davis_CoverLetter_', '').replace('.txt', '')
        
        # Convert underscore naming to proper company names
        company_mapping = {
            'Blizzard_Quality_Manager': 'Blizzard Entertainment',
            'LightWonder_Producer': 'Light & Wonder',
            'Mandolin_QA_Manager': 'Mandolin',
            'Agtonomy_QA_Manager': 'Agtonomy',
            'Activision_Director_QA': 'Activision',
            'Riot_Games_TFT_QA_Manager': 'Riot Games',
            'Netflix_Consumer_Insights_Manager': 'Netflix',
            'Infantino_Compliance_QA_Manager': 'Infantino',
            'QualStaff_Sr_Manager_GxP': 'QualStaff Resources',
            'Caesars_QA_Manager_Corporate': 'Caesars Entertainment'
        }
        
        return company_mapping.get(name, name.replace('_', ' '))
    
    def convert_all_letters(self):
        """Convert all cover letter text files to Word documents."""
        print("🚀 Cover Letter to Word Document Converter")
        print("=" * 50)
        
        # Find all cover letter text files
        pattern = "Kenneth_Davis_CoverLetter_*.txt"
        letter_files = glob.glob(pattern)
        
        if not letter_files:
            print("❌ No cover letter files found matching pattern: " + pattern)
            return False
        
        print(f"✅ Found {len(letter_files)} cover letter files")
        
        converted_count = 0
        
        for txt_file in letter_files:
            try:
                # Read the text file
                print(f"\n📄 Processing: {txt_file}")
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Extract company name for proper formatting
                company_name = self.extract_company_name(os.path.basename(txt_file))
                print(f"   Company: {company_name}")
                
                # Create Word document
                doc = self.create_formatted_document(content, company_name)
                
                # Generate output filename
                output_filename = f"Kenneth_Davis_Cover_Letter_{company_name.replace(' ', '_').replace('&', 'and')}.docx"
                output_path = os.path.join(self.desktop_path, output_filename)
                
                # Save the document
                doc.save(output_path)
                print(f"   ✅ Saved: {output_path}")
                converted_count += 1
                
            except Exception as e:
                print(f"   ❌ Error processing {txt_file}: {str(e)}")
                continue
        
        print(f"\n🎉 Conversion complete! {converted_count}/{len(letter_files)} files converted")
        print(f"📁 Output location: {self.desktop_path}")
        return converted_count > 0
    
    def list_created_files(self):
        """List the created Word documents."""
        pattern = os.path.join(self.desktop_path, "Kenneth_Davis_Cover_Letter_*.docx")
        docx_files = glob.glob(pattern)
        
        if docx_files:
            print("\n📋 Created Word Documents:")
            for docx_file in sorted(docx_files):
                filename = os.path.basename(docx_file)
                print(f"   • {filename}")
        else:
            print("\n❌ No Word documents found on Desktop")

def main():
    try:
        # Check if python-docx is available
        import docx
    except ImportError:
        print("❌ python-docx library not found!")
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call(['pip', 'install', 'python-docx'])
        print("✅ python-docx installed successfully")
    
    converter = CoverLetterConverter()
    success = converter.convert_all_letters()
    
    if success:
        converter.list_created_files()
        print("\n✅ SUCCESS: All cover letters converted to Word documents!")
        print("📁 Check your Desktop for the .docx files")
    else:
        print("\n❌ FAILED: Could not convert cover letters")
    
    return 0 if success else 1

if __name__ == "__main__":
    exit(main())
